import os
from datetime import datetime
import pdfplumber
import docx


def list_files(directory: str, extension: str = None) -> dict:
    try:
        if not os.path.exists(directory):
            return {"status": "error", "error": "Directory does not exist"}
        files = []
        for file in os.listdir(directory):
            path = os.path.join(directory, file)

            if os.path.isfile(path):
                if extension and not file.lower().endswith(extension.lower()):
                    continue

                stat = os.stat(path)

                files.append({
                    "name": file,
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })

        return {"status": "success", "data": files}

    except Exception as e:
        return {"status": "error", "error": str(e)}


def read_file(file_path: str) -> dict:
    try:
        if not os.path.exists(file_path):
            return {"status": "error", "error": "File not found"}

        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

        elif file_path.endswith(".pdf"):
            content = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    content += page.extract_text() or ""

        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)

            content = "\n".join(content)

        else:
            return {"status": "error", "error": "Unsupported file type"}

        return {
            "status": "success",
            "data": {"content": content}
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}


def search_in_file(file_path: str, keyword: str) -> dict:
    try:
        res = read_file(file_path)
        if res["status"] == "error":
            return res

        lines = res["data"]["content"].split("\n")
        matches = []

        for i, line in enumerate(lines):
            if keyword.lower() in line.lower():
                context = " ".join(lines[max(0, i-1):min(len(lines), i+2)])
                matches.append({
                    "line_number": i + 1,
                    "match": line.strip(),
                    "context": context
                })

        return {"status": "success", "data": matches}

    except Exception as e:
        return {"status": "error", "error": str(e)}


def write_file(file_path: str, content: str) -> dict:
    try:
        if not content.strip():
            return {"status": "error", "error": "Empty content"}

        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        return {"status": "success", "path": file_path}

    except Exception as e:
        return {"status": "error", "error": str(e)}
    


def generate_summary_filename(original_path: str) -> str:
    base_name = os.path.basename(original_path)
    name_without_ext = os.path.splitext(base_name)[0]

    clean_name = name_without_ext.replace(" ", "_").lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    return f"summary_{clean_name}_{timestamp}.txt"
